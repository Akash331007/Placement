import io
from pypdf import PdfReader
from docx import Document

def extract_text_from_pdf(file_bytes: bytes) -> str:
    pdf_file = io.BytesIO(file_bytes)
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    docx_file = io.BytesIO(file_bytes)
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text

def extract_text(file_bytes: bytes, file_type: str) -> str:
    file_type = file_type.lower().strip('.')
    if file_type == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif file_type in ["docx", "doc"]:
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
